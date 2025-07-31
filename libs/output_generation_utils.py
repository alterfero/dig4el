import json
from docx import Document
from libs import knowledge_graph_utils as kgu
from io import BytesIO
import re
from collections import defaultdict
import pandas as pd
import graphviz
from html import escape

import jinja2

def generate_transcription_doc(cq, target_language, pivot_language):
    if target_language in ["English", ""]:
        target_language = "target language"

    document = Document()
    document.add_heading('Conversational Questionnaire', 0)
    document.add_heading(f'"{cq["title"]}"', 0)
    document.add_heading("Information", 1)
    table = document.add_table(rows=5, cols=2, style="Light Grid")  # Fixed column count
    row_cells0 = table.rows[0].cells
    row_cells0[0].text = "Target language"
    if target_language != "target language":
        row_cells0[1].text = target_language
    row_cells1 = table.rows[1].cells
    row_cells1[0].text = "Transcription made by"
    row_cells2 = table.rows[2].cells
    row_cells2[0].text = f"Content in {target_language} provided by"
    row_cells3 = table.rows[3].cells
    row_cells3[0].text = "Location"
    row_cells4 = table.rows[4].cells
    row_cells4[0].text = "Date"

    document.add_page_break()

    document.add_heading("Full dialog in English", 1)
    document.add_paragraph(cq["context"])
    dialog_table = document.add_table(rows=1, cols=3, style="Table Grid")  # Fixed column count
    hdr_cells = dialog_table.rows[0].cells
    hdr_cells[0].text = "Index"
    hdr_cells[1].text = cq["speakers"]["A"]["name"]
    hdr_cells[2].text = cq["speakers"]["B"]["name"]
    dialog_length = len(cq["dialog"])
    for index in range(1, dialog_length + 1):
        content = cq["dialog"][str(index)]
        row_cells = dialog_table.add_row().cells
        if content.get("legacy index", "") != "":
            i = content["legacy index"]
        else:
            i = str(index)
        row_cells[0].text = i
        if content["speaker"] == "A":
            row_cells[1].text = content["text"]
            row_cells[2].text = ""
        elif content["speaker"] == "B":
            row_cells[1].text = ""
            row_cells[2].text = content["text"]

    document.add_page_break()

    document.add_heading("Transcription", 1)

    for index in range(1, dialog_length + 1):
        content = cq["dialog"][str(index)]
        if content.get("legacy index", "") != "":
            i = content["legacy index"]
        else:
            i = str(index)
        document.add_paragraph("")
        document.add_heading(f'{i}: {content["text"]}', 2)
        transcription_table = document.add_table(rows=0, cols=2, style="Table Grid")
        row_cells = transcription_table.add_row().cells
        row_cells[0].text = "Pivot (if not English)"
        row_cells[1].text = ""
        row_cells = transcription_table.add_row().cells
        row_cells[0].text = f'Sentence in {target_language}'
        row_cells[1].text = ""
        document.add_paragraph(" ")
        p = document.add_paragraph()
        p.add_run("Connections between word(s) and concept(s):").bold = True
        document.add_paragraph("In this segment, which word(s) contribute to which concept(s) below? One word can appear in multiple concepts, or in none. Multiple words can appear in a single concept.")
        concept_table = document.add_table(rows=1, cols=2, style="Table Grid")
        hdr_cells = concept_table.rows[0].cells
        hdr_cells[0].text = "Concept"
        hdr_cells[1].text = "Word(s) contributing to this concept"
        if content.get("intent", []) != []:
            row_cells = concept_table.add_row().cells
            row_cells[0].text = f'Intent: {"+".join(content["intent"])}'
            row_cells[1].text = ""
        if content.get("predicate", []) != []:
            row_cells = concept_table.add_row().cells
            row_cells[0].text = f'Type of predicate: {"+".join(content["predicate"])}'
            row_cells[1].text = ""
        for c in content["concept"]:
            row_cells = concept_table.add_row().cells
            row_cells[0].text = c
            row_cells[1].text = ""

    # Save to a BytesIO buffer instead of disk
    docx_buffer = BytesIO()
    document.save(docx_buffer)
    docx_buffer.seek(0)  # Reset buffer position

    return docx_buffer



def generate_docx_from_kg_index_list(kg, delimiters, kg_index_list):
    document = Document()
    document.add_heading('Partial corpus', 0)
    item_counter = 0
    for index in kg_index_list:
        item_counter += 1
        data = kg[index]
        gloss = kgu.build_super_gloss_df(kg, index, delimiters,  output_dict=True)
        document.add_paragraph("""Entry {}""".format(item_counter), style='List Bullet')
        pex = document.add_paragraph(" ", "Normal")
        pex.add_run(f'{data["recording_data"]["translation"]}', style='Strong')
        document.add_paragraph(f'{data["sentence_data"]["text"]}', "Normal")
        document.add_paragraph(f'''
        Intent: {", ".join(data["sentence_data"]["intent"])}
        Type of predicate: {", ".join(data["sentence_data"]["predicate"])}
        ''', "Normal")

        # Define the number of rows (one header + words) and columns
        num_words = len(gloss)
        table = document.add_table(rows=num_words + 1, cols=4, style="Light Shading Accent 1")  # Fixed column count

        # Populate header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ""
        hdr_cells[1].text = "Concept"
        hdr_cells[2].text = "Internal Particularisation"
        hdr_cells[3].text = "Relational Particularisation"

        # Populate the table with data
        for i, w in enumerate(gloss, start=1):
            row_cells = table.rows[i].cells
            row_cells[0].text = w["word"]
            row_cells[1].text = w["concept"]
            row_cells[2].text = w["internal particularization"]
            row_cells[3].text = w["relational particularization"]
        document.add_paragraph("""   """)

    # Save to a BytesIO buffer instead of disk
    docx_buffer = BytesIO()
    document.save(docx_buffer)
    docx_buffer.seek(0)  # Reset buffer position

    return docx_buffer

_MAX_WORD_COLS = 63             # Word hard‑limit on columns  :contentReference[oaicite:0]{index=0}
_BAD_XML_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")

def _safe_text(value: object) -> str:
    """Return XML‑safe, non‑None, length‑capped string for a table cell."""
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return ""
    text = str(value)
    # strip control chars Word’s XML parser can’t handle
    text = _BAD_XML_CHARS_RE.sub("", text)
    # Word has no documented per‑cell limit, but 32 K keeps you well clear
    return text[:32760]

def add_dataframe_table(doc: Document,
                        df: pd.DataFrame,
                        style: str = "LightShading-Accent1") -> None:
    """
    Render *df* as a Word table and append it to *doc*, coping with:
      • empty frames
      • >63 columns
      • missing / localised style names
      • NaN / None / weird unicode in the data
    """
    if df.empty or df.shape[1] == 0:
        return                                      # nothing to do

    # Word refuses tables wider than 63 columns
    if df.shape[1] > _MAX_WORD_COLS:
        df = df.iloc[:, :_MAX_WORD_COLS]

    # build the table (first row = header row)
    table = doc.add_table(rows=1, cols=df.shape[1])

    # try the requested style first, fall back gracefully if Word/doc lacks it
    try:
        # API style names omit the hyphen: “Light Shading – Accent 1”
        # becomes “Light Shading Accent 1” :contentReference[oaicite:1]{index=1}
        table.style = style
    except KeyError:
        pass                                         # default style is OK

    # ---- header row -----------------------------------------------------
    for cell, col in zip(table.rows[0].cells, df.columns):
        cell.text = _safe_text(col)

    # ---- data rows ------------------------------------------------------
    for _, row in df.iterrows():
        for cell, val in zip(table.add_row().cells, row):
            cell.text = _safe_text(val)


STRIP_IDX_RE = re.compile(r'(_\d+|\(\d+\))$', re.VERBOSE)  # to remove indexes _1 or (1) added to target words
NONWORD_RE = re.compile(r'\W+')
WORD_COL_CANDIDATES = ("target_word", "word")

STRIP_IDX_RE = re.compile(r"(_\d+|\(\d+\))$")      # strips _2  or (2)


# ------------------------------------------------------------------
#  main helper
# ------------------------------------------------------------------
def _prepare_words(df: pd.DataFrame) -> pd.DataFrame:
    # 1) find the column that holds the word
    for col in WORD_COL_CANDIDATES:
        if col in df.columns:
            word_col = col
            break
    else:
        raise KeyError(
            f"None of {WORD_COL_CANDIDATES} found in DataFrame columns "
            f"{list(df.columns)}"
        )

    out = df.copy()

    # ------------------------------------------------------------------
    # Strip “_2” or “(2)” only once and reuse the result everywhere
    # ------------------------------------------------------------------
    clean_word = (
        out[word_col]
        .fillna("")
        .astype(str)
        .str.replace(r"(_\d+|\(\d+\))$", "", regex=True)
    )

    # ---- what the reader will see ------------------------------------
    out["display_word"] = clean_word

    # ---- Graphviz node‑id (slug of the cleaned word) ------------------
    out["word_node_id"] = (
        "word_" + clean_word.str.replace(r"\W+", "_", regex=True)
    )

    # ---- concept node‑id (unchanged) ----------------------------------
    out["concept_node_id"] = (
        "concept_" +
        (out["concept"].astype(str)
         + "|" + out["IP"].astype(str)
         + "|" + out["RP"].astype(str))
        .apply(hash).astype(str)
    )

    return out

# ------------------------------------------------------------------
#  main helper
# ------------------------------------------------------------------
def add_concept_graph(doc: Document, df: pd.DataFrame) -> None:
    """
    Build a left‑to‑right bipartite graph:
        target_word  ──▶  concept (IP, RP)
    and insert it as a picture into *doc*.

    Required df columns: target_word, concept, IP, RP
    """
    df = _prepare_words(df)          # ← use the cleaned copy

    g = graphviz.Digraph('WordConcept', format='png')
    g.attr(rankdir='LR', splines='true', overlap='false')

    # ---------- left column: target words ---------------------------
    with g.subgraph(name='cluster_targets') as s:
        s.attr(rank='same')
        for node_id, label in (df[['word_node_id', 'display_word']]
                               .drop_duplicates()
                               .sort_values('display_word')
                               .itertuples(index=False, name=None)):
            s.node(node_id, label,
                   shape='box', style='rounded,filled',
                   fillcolor='#f0f8ff', fontsize='14')

    # ---------- right column: concept + IP + RP ---------------------
    with g.subgraph(name='cluster_concepts') as s:
        s.attr(rank='same')
        for node_id, c, ip, rp in (
                df[["concept_node_id", "concept", "IP", "RP"]]
                        .drop_duplicates()
                        .sort_values(["concept", "IP", "RP"])
                        .itertuples(index=False, name=None)
        ):
            label = f"{c}\nIP={ip}  RP={rp}"
            s.node(node_id, label,
                   shape="ellipse", style="filled",
                   fillcolor="#ffffe0", fontsize="12")

    # ---------- edges -----------------------------------------------
    for _, r in df.iterrows():
        g.edge(r.word_node_id, r.concept_node_id)

    # ---------- render & insert -------------------------------------
    stream = BytesIO(g.pipe(format='png'))
    doc.add_picture(stream, width=doc.sections[0].page_width * 0.6)
    doc.add_paragraph()           # blank line after the picture


def generate_docx_from_hybrid_output(content, language, gloss_format="table"):
    document = Document()
    document.add_heading('Learning {}'.format(language), 0)
    if "introduction" in content.keys():
        document.add_heading('Introduction', 1)
        document.add_paragraph(content["introduction"]["description"])
    for chapter in content["chapters"]:
        document.add_heading(chapter["title"], 1)
        document.add_paragraph(chapter["explanation"])
        e_counter = 0
        for example in chapter["examples"]:
            e_counter += 1
            document.add_paragraph("""Example {}""".format(e_counter), style='List Bullet')
            pex = document.add_paragraph(" ", "Normal")
            pex.add_run(f'{example["target"]}', style='Strong')
            document.add_paragraph(f'{example["english"]}', "Normal")

            df = parse_alterlingua(example["gloss"])

            if gloss_format == "table":
                add_dataframe_table(document, df)
            elif gloss_format == "graph":
                add_concept_graph(document, df)

            document.add_paragraph("""   """)

    # Save to a BytesIO buffer instead of disk
    docx_buffer = BytesIO()
    document.save(docx_buffer)
    docx_buffer.seek(0)  # Reset buffer position

    return docx_buffer

def generate_plain_language_docx_from_hybrid_output(content, language):
    document = Document()
    document.add_heading('Learning {}'.format(language), 0)
    if "introduction" in content.keys():
        document.add_heading('Introduction', 1)
        document.add_paragraph(content["introduction"]["description"])
    for chapter in content["chapters"]:
        document.add_heading(chapter["title"], 1)
        document.add_paragraph(chapter["explanation"])
        e_counter = 0
        for example in chapter["examples"]:
            e_counter += 1
            document.add_paragraph("""Example {}""".format(e_counter), style='List Bullet')
            pex = document.add_paragraph(" ", "Normal")
            pex.add_run(f'{example["target"]}', style='Strong')
            document.add_paragraph(f'{example["english"]}', "Normal")
            document.add_paragraph("""   """)
            document.add_paragraph(f'{example["gloss"]}', "Normal")
            document.add_paragraph("""   """)

    # Save to a BytesIO buffer instead of disk
    docx_buffer = BytesIO()
    document.save(docx_buffer)
    docx_buffer.seek(0)  # Reset buffer position

    return docx_buffer


def generate_lesson_docx_from_aggregated_output(content, indi_language, readers_language):
    loc = {
        "English": {
            "introduction": "Introduction",
            "for example": "For example",
            "in conclusion": "In Conclusion",
            "more_examples": "More examples",
            "learning_the_x_language_grammar": "Learning the X language grammar"
        },
        "Bislama": {
            "introduction": "Introdaksen",
            "for example": "Foa eksampol",
            "in conclusion": "Long fin",
            "more_examples": "Mo eksampol",
            "learning_the_x_language_grammar": "Lernim graema blong lanwis X"
        },
        "French": {
            "introduction": "Introduction",
            "for example": "Par exemple",
            "in conclusion": "En conclusion",
            "more_examples": "Plus d'exemples",
            "learning_the_x_language_grammar": "Apprentissage de la grammaire de la langue X"
        },
        "Japanese": {
            "introduction": "はじめに",
            "for example": "例えば",
            "in conclusion": "結論として",
            "more_examples": "さらに例",
            "learning_the_x_language_grammar": "X言語の文法を学ぶ"
        },
        "Swedish": {
            "introduction": "Introduktion",
            "for example": "Till exempel",
            "in conclusion": "Sammanfattningsvis",
            "more_examples": "Fler exempel",
            "learning_the_x_language_grammar": "Att lära sig X-språkets grammatik"
        }
    }
    locit = loc.get(readers_language, "English")
    learning_X = locit["learning_the_x_language_grammar"].replace("X", indi_language)

    document = Document()

    document.add_heading(content["title"], level=1)
    document.add_paragraph("  ")
    document.add_paragraph(content["introduction"])

    for section in content["sections"]:
        document.add_heading(section["focus"], level=2)
        document.add_paragraph(section["description"]["description"])
        document.add_paragraph(locit["for example"] + ": ")
        pex = document.add_paragraph(" ", "Normal")
        pex.add_run(f'{section["example"]["target_sentence"]}', style='Strong')
        document.add_paragraph(f'({section["example"]["source_sentence"]})', "Normal")
        document.add_paragraph(section["example"]["description"])

    document.add_heading(locit["in conclusion"], level=1)
    document.add_paragraph(" ")
    document.add_paragraph(content["conclusion"])

    document.add_heading(locit["more_examples"] + ": ")
    document.add_paragraph(" ")
    for i, s in enumerate(content["translation_drills"]):
        aex = document.add_paragraph(" ", style="List Bullet")
        aex.add_run(f'{s["target"]}', style='Strong')
        document.add_paragraph(f'({s["source"]})', "Normal")

    # Save to a BytesIO buffer instead of disk
    docx_buffer = BytesIO()
    document.save(docx_buffer)
    docx_buffer.seek(0)  # Reset buffer position

    return docx_buffer




def generate_docx_from_grammar_json(grammar_json, language):
    h1_index = 0
    h2_index = 0
    document = Document()
    document.add_heading('{}: Elements of grammar.'.format(language), 0)
    for topic, content in grammar_json.items():
        h1_index += 1
        document.add_heading(f"{str(h1_index)}. {topic}", 1)
        for parameter, description in content.items():
            h2_index += 1
            document.add_heading(f"{str(h1_index)}.{str(h2_index)}. {topic}: {parameter}", 2)
            intro_text = f"""
            In {language}, {parameter} is mainly {description["main value"]}
            """
            p = document.add_paragraph(intro_text)
            for value, examples in description["examples by value"].items():
                document.add_heading(f"""Example of {value}: """, 3)
                example_counter = 0
                for example in examples:
                    example_counter += 1
                    document.add_paragraph("""Example {}""".format(example_counter), style='List Bullet')
                    pex = document.add_paragraph(" ", "Normal")
                    pex.add_run(f'{example["translation"]}', style='Strong')
                    document.add_paragraph(f'{example["english sentence"]}', "Normal")

                    # Define the number of rows (one header + words) and columns
                    num_words = len(example["gloss"])
                    table = document.add_table(rows=num_words + 1, cols=4, style="Light Shading Accent 1")  # Fixed column count

                    # Populate header row
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = language
                    hdr_cells[1].text = "Concept"
                    hdr_cells[2].text = "Internal Particularisation"
                    hdr_cells[3].text = "Relational Particularisation"

                    # Populate the table with data
                    for i, (w, g) in enumerate(example["gloss"].items(), start=1):
                        row_cells = table.rows[i].cells
                        row_cells[0].text = w
                        row_cells[1].text = g["concept"]
                        row_cells[2].text = g["internal particularization"]
                        row_cells[3].text = g["relational particularization"]
                    document.add_paragraph("""   """)

        # Save to a BytesIO buffer instead of disk
        docx_buffer = BytesIO()
        document.save(docx_buffer)
        docx_buffer.seek(0)  # Reset buffer position
        return docx_buffer
    return None


def parse_alterlingua(text):
    # Regex to find all target word blocks
    word_blocks = re.findall(r'(\S+)<([^>]*)>', text)

    results = []

    for target_word, content in word_blocks:
        if not content.strip():
            continue  # skip empty concept blocks

        # Split multiple concepts joined by '&'
        concepts = [c.strip() for c in content.split('&')]

        for idx, concept in enumerate(concepts, start=1):
            # Regex to extract the concept name and its IP / RP fields
            concept_match = re.match(r'([^()]+)\(([^)]*)\)', concept)
            if concept_match:
                concept_name, fields = concept_match.groups()
                concept_name = concept_name.strip()

                # Extract IP and RP using regex or string methods
                ip_match = re.search(r'IP:\s*([^|]*)', fields)
                rp_match = re.search(r'RP:\s*([^|]*)', fields)
                ip = ip_match.group(1).strip() if ip_match else ""
                rp = rp_match.group(1).strip() if rp_match else ""

                entry = {
                    "target_word": f"{target_word}({idx})" if len(concepts) > 1 else target_word,
                    "concept": concept_name,
                    "IP": ip,
                    "RP": rp
                }
                results.append(entry)

    return pd.DataFrame(results)

